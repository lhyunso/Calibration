"""
DOCX calibration report generator.
Design inspired by 'Calibration sheet template.docx'.
"""
import os
import io
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np

from processing.calibration import ChannelCalibration
from sensors.base import SensorConfig


# ── Colour palette ──────────────────────────────────────────────────────────
C_HEADER_BG  = RGBColor(0x1F, 0x49, 0x7D)   # dark navy
C_HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)   # white
C_SECTION_BG = RGBColor(0xD6, 0xE4, 0xF0)   # light blue
C_ROW_ALT    = RGBColor(0xF2, 0xF7, 0xFB)   # very light blue
C_PASS       = RGBColor(0x00, 0x70, 0xC0)   # blue (within tolerance)
C_FAIL       = RGBColor(0xC0, 0x00, 0x00)   # red  (out of tolerance)
C_BORDER     = RGBColor(0xA0, 0xA0, 0xA0)   # grey


# ── Helper utilities ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(rgb))
    tcPr.append(shd)


def _set_cell_border(cell, sides=("top", "bottom", "left", "right"), size=4, color="A0A0A0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para(cell, text: str, bold=False, italic=False, size=9,
          align=WD_ALIGN_PARAGRAPH.LEFT, color: Optional[RGBColor] = None,
          font_name="맑은 고딕"):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.color.rgb = color if color else RGBColor(0, 0, 0)
    return p


def _header_cell(cell, text: str, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    _set_cell_bg(cell, C_HEADER_BG)
    _para(cell, text, bold=True, size=size, align=align, color=C_HEADER_FG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _section_cell(cell, text: str, size=9):
    _set_cell_bg(cell, C_SECTION_BG)
    _para(cell, text, bold=True, size=size, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fmt(val: Optional[float], digits=6) -> str:
    if val is None:
        return "-"
    return f"{val:.{digits}f}"


def _pass_color(dev: float, tolerance: float) -> RGBColor:
    return C_PASS if abs(dev) <= tolerance else C_FAIL


# ── Chart generation ─────────────────────────────────────────────────────────

def _make_gain_chart(cal: ChannelCalibration) -> io.BytesIO:
    """
    Create a scatter+line chart showing:
      - Measured voltage vs Reference voltage (raw)
      - Fitted line (gain)
    Returns PNG bytes in a BytesIO buffer.
    """
    rs = sorted(cal.voltages_avg.keys())
    v_meas = [cal.voltages_avg[r] for r in rs]
    v_ref  = [cal.voltage_ref[r]  for r in rs]

    fig, ax = plt.subplots(figsize=(5.5, 3.2))
    ax.scatter(v_ref, v_meas, color="#1F497D", zorder=5, s=50, label="Measured")

    # Fit line
    if len(v_ref) >= 2:
        coeffs = np.polyfit(v_ref, v_meas, 1)
        x_line = np.linspace(min(v_ref)*1.05, max(v_ref)*1.05, 100)
        y_line = np.polyval(coeffs, x_line)
        ax.plot(x_line, y_line, color="#C00000", linewidth=1.5,
                label=f"Gain={cal.gain:.6f}")

    ax.set_xlabel("Reference Voltage (V)", fontsize=8)
    ax.set_ylabel("Measured Voltage (V)", fontsize=8)
    ax.set_title(f"Gain Calibration — {cal.channel}", fontsize=9, fontweight="bold")
    ax.legend(fontsize=7)
    ax.grid(True, linestyle="--", alpha=0.4)
    ax.tick_params(labelsize=7)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_deviation_chart(cal: ChannelCalibration, resistances: List[float]) -> io.BytesIO:
    """
    Bar chart showing deviation after gain correction for both offset methods,
    with tolerance bands.
    """
    rs = sorted([r for r in resistances if r in cal.dev_final_100])
    dev_100  = [cal.dev_final_100.get(r, 0) for r in rs]
    dev_mean = [cal.dev_final_mean.get(r, 0) for r in rs]
    tol      = cal.tolerance_max_100.get(rs[0], 0.385) if rs else 0.385

    x = np.arange(len(rs))
    width = 0.35

    fig, ax = plt.subplots(figsize=(5.5, 3.2))
    bars1 = ax.bar(x - width/2, dev_100,  width, label="Method 2-1 (100Ω offset)",
                   color="#1F497D", alpha=0.8)
    bars2 = ax.bar(x + width/2, dev_mean, width, label="Method 2-2 (Mean offset)",
                   color="#ED7D31", alpha=0.8)

    ax.axhline(y= tol, color="red",  linestyle="--", linewidth=1.0, label=f"+Tol ({tol}Ω)")
    ax.axhline(y=-tol, color="blue", linestyle="--", linewidth=1.0, label=f"-Tol ({-tol}Ω)")
    ax.axhline(y=0,    color="black",linestyle="-",  linewidth=0.5)

    ax.set_xticks(x)
    ax.set_xticklabels([f"{int(r)}Ω" for r in rs], fontsize=7)
    ax.set_ylabel("Deviation (Ω)", fontsize=8)
    ax.set_title(f"Calibrated Deviation — {cal.channel}", fontsize=9, fontweight="bold")
    ax.legend(fontsize=6, loc="upper right")
    ax.grid(True, axis="y", linestyle="--", alpha=0.4)
    ax.tick_params(labelsize=7)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


# ── Document builder ──────────────────────────────────────────────────────────

class CalibrationDocxWriter:
    """Generates a calibration report DOCX."""

    def __init__(
        self,
        sensor: SensorConfig,
        calibrations: Dict[str, ChannelCalibration],
        metadata: Optional[Dict[str, str]] = None,
    ):
        self.sensor = sensor
        self.calibrations = calibrations
        self.channels = sorted(calibrations.keys())
        self.meta = metadata or {}
        self.resistances: List[float] = []
        if calibrations:
            first = next(iter(calibrations.values()))
            self.resistances = sorted(first.voltages_avg.keys())

    def _page_setup(self, doc: Document):
        section = doc.sections[0]
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ── Cover / Summary page ─────────────────────────────────────────────────

    def _add_cover(self, doc: Document):
        """Calibration report cover page."""
        # ── Document title ────────────────────────────────────────────────────
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run("CALIBRATION REPORT")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "맑은 고딕"
        run.font.color.rgb = C_HEADER_BG

        # Doc number / revision line
        dn_para = doc.add_paragraph()
        dn_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc_no = self.meta.get("doc_number", f"CAL-{datetime.now().strftime('%Y')}-0001")
        rev    = self.meta.get("revision", "00")
        dn_run = dn_para.add_run(f"문서번호: {doc_no}   Rev: {rev}")
        dn_run.font.size = Pt(8)
        dn_run.font.name = "맑은 고딕"
        dn_run.font.color.rgb = C_HEADER_BG

        doc.add_paragraph()  # spacer

        # ── Info table (2-column label/value layout) ──────────────────────────
        tbl = doc.add_table(rows=14, cols=4)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_widths = [Cm(4.0), Cm(5.0), Cm(4.0), Cm(5.0)]
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

        def _info(row_idx, label_l, val_l, label_r="", val_r=""):
            row = tbl.rows[row_idx]
            _section_cell(row.cells[0], label_l)
            _para(row.cells[1], val_l, size=9)
            if label_r:
                _section_cell(row.cells[2], label_r)
                _para(row.cells[3], val_r, size=9)
            else:
                _para(row.cells[2], "", size=9)
                _para(row.cells[3], "", size=9)

        # Section 1: Module Info
        _info(0,  "[ 모듈 정보 ]",           "",
                   "",                         "")
        _set_cell_bg(tbl.rows[0].cells[0], C_HEADER_BG)
        _para(tbl.rows[0].cells[0], "[ 모듈 정보 ]",
              bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
        tbl.rows[0].cells[0].merge(tbl.rows[0].cells[3])

        _info(1,  "모듈명 (Module Name)",     self.meta.get("module_name", ""),
                   "모델명 (Model No.)",       self.meta.get("model", ""))
        _info(2,  "시리얼 번호 (S/N)",         self.meta.get("serial", ""),
                   "제조사 (Manufacturer)",    self.meta.get("manufacturer", ""))
        _info(3,  "FW / SW 버전",             self.meta.get("fw_version", ""),
                   "",                         "")

        # Section 2: Calibration Conditions
        _set_cell_bg(tbl.rows[4].cells[0], C_HEADER_BG)
        _para(tbl.rows[4].cells[0], "[ 교정 조건 ]",
              bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
        tbl.rows[4].cells[0].merge(tbl.rows[4].cells[3])

        _info(5,  "교정 일자 (Date)",          self.meta.get("date", datetime.now().strftime("%Y.%m.%d")),
                   "교정 장소 (Location)",     self.meta.get("location", ""))
        _info(6,  "담당자 (Technician)",        self.meta.get("operator", ""),
                   "온도 / 습도",              self.meta.get("temp_humidity", ""))

        # Section 3: Calibration Setup
        _set_cell_bg(tbl.rows[7].cells[0], C_HEADER_BG)
        _para(tbl.rows[7].cells[0], "[ 교정 설정 ]",
              bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
        tbl.rows[7].cells[0].merge(tbl.rows[7].cells[3])

        _info(8,  "센서 타입 (Sensor)",        self.sensor.name,
                   "케이블 (Cable)",           self.meta.get("cable", "전용케이블"))
        _info(9,  "공칭 저항 (Nominal R)",     f"{self.sensor.r_nominal:.0f} Ω",
                   "Inst. Amp. Gain",          str(self.meta.get("inst_amp_gain", "1")))
        _info(10, "모사저항 값",
                  " / ".join(f"{r:.0f}Ω" for r in self.resistances),
                   "허용 편차 (Tolerance)",   f"±{self.sensor.tolerance_ohm:.3f} Ω")
        _info(11, "샘플링 속도",
                  f"{self.meta.get('sampling_hz', 100)} Hz",
                   "측정 시간",               f"{self.meta.get('duration_sec', '-')} 초")
        _info(12, "채널 수 (Channels)",        str(len(self.channels)),
                   "레퍼런스 수식",            self.sensor.ref_formula)
        _info(13, "비고 (Remarks)",            self.meta.get("remarks", ""),
                   "",                         "")

        for row in tbl.rows:
            for cell in row.cells:
                _set_cell_border(cell)

    # ── Channel summary table ─────────────────────────────────────────────────

    def _add_summary(self, doc: Document):
        doc.add_paragraph()
        h = doc.add_paragraph("2-1  Calibration Summary (100Ω Offset Method)")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(11)
        h.runs[0].font.color.rgb = C_HEADER_BG

        rs = self.resistances
        n_rs = len(rs)
        # cols: Channel | Excitation | Gain | Offset | dev×n
        # A4 usable width ≈ 17cm; distribute proportionally
        n_cols = 4 + n_rs
        tbl = doc.add_table(rows=2 + len(self.channels), cols=n_cols)
        tbl.style = "Table Grid"
        tbl.autofit = False
        _col_w = [Cm(3.0), Cm(2.2), Cm(2.4), Cm(2.4)] + [Cm(max(1.4, (17.0 - 10.0) / n_rs))] * n_rs
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                cell.width = _col_w[i]

        # Header row 1
        row0 = tbl.rows[0]
        for i, label in enumerate(["채널", "Excitation", "Gain", "Offset(R_nom)"]):
            _header_cell(row0.cells[i], label)
        for j, r in enumerate(rs):
            _header_cell(row0.cells[4 + j], f"{r:.0f}Ω")

        # Header row 2: units
        row1 = tbl.rows[1]
        for i, u in enumerate(["-", "V/Ω", "-", "Ω"] + ["dev(Ω)"] * n_rs):
            _header_cell(row1.cells[i], u, size=7)

        # Data rows
        for ri, ch in enumerate(self.channels):
            cal = self.calibrations[ch]
            row = tbl.rows[2 + ri]
            bg = C_ROW_ALT if ri % 2 == 0 else None

            cells = [
                ch,
                f"{cal.excitation:.4f}",
                f"{cal.gain:.6f}",
                f"{cal.offset_100:.6f}",
            ]
            for i, txt in enumerate(cells):
                _para(row.cells[i], txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                if bg:
                    _set_cell_bg(row.cells[i], bg)

            for j, r in enumerate(rs):
                dev = cal.dev_final_100.get(r)
                tol = cal.tolerance_max_100.get(r, self.sensor.tolerance_ohm)
                txt = _fmt(dev, 4) if dev is not None else "-"
                c = row.cells[4 + j]
                _para(c, txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
                      color=_pass_color(dev, tol) if dev is not None else None)
                if bg:
                    _set_cell_bg(c, bg)

        for r in tbl.rows:
            for cell in r.cells:
                _set_cell_border(cell)

        # ── Method 2-2 summary ──
        doc.add_paragraph()
        h2 = doc.add_paragraph("2-2  Calibration Summary (Mean Offset Method)")
        h2.runs[0].bold = True
        h2.runs[0].font.size = Pt(11)
        h2.runs[0].font.color.rgb = C_HEADER_BG

        tbl2 = doc.add_table(rows=2 + len(self.channels), cols=n_cols)
        tbl2.style = "Table Grid"
        tbl2.autofit = False
        for row in tbl2.rows:
            for i, cell in enumerate(row.cells):
                cell.width = _col_w[i]
        row0b = tbl2.rows[0]
        for i, label in enumerate(["채널", "Excitation", "Gain", "Offset(Mean)"]):
            _header_cell(row0b.cells[i], label)
        for j, r in enumerate(rs):
            _header_cell(row0b.cells[4 + j], f"{r:.0f}Ω")
        row1b = tbl2.rows[1]
        for i, u in enumerate(["-", "V/Ω", "-", "Ω"] + ["dev(Ω)"] * n_rs):
            _header_cell(row1b.cells[i], u, size=7)

        for ri, ch in enumerate(self.channels):
            cal = self.calibrations[ch]
            row = tbl2.rows[2 + ri]
            bg = C_ROW_ALT if ri % 2 == 0 else None
            cells = [ch, f"{cal.excitation:.4f}", f"{cal.gain:.6f}", f"{cal.offset_mean:.6f}"]
            for i, txt in enumerate(cells):
                _para(row.cells[i], txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                if bg:
                    _set_cell_bg(row.cells[i], bg)
            for j, r in enumerate(rs):
                dev = cal.dev_final_mean.get(r)
                tol = cal.tolerance_max_mean.get(r, self.sensor.tolerance_ohm)
                txt = _fmt(dev, 4) if dev is not None else "-"
                c = row.cells[4 + j]
                _para(c, txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
                      color=_pass_color(dev, tol) if dev is not None else None)
                if bg:
                    _set_cell_bg(c, bg)

        for r in tbl2.rows:
            for cell in r.cells:
                _set_cell_border(cell)

    # ── Per-channel detail page ──────────────────────────────────────────────

    def _add_channel_page(self, doc: Document, ch_idx: int, ch: str):
        doc.add_page_break()
        cal = self.calibrations[ch]
        rs = self.resistances

        # ---- Page header ----
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.add_run(f"Calibration Sheet ({ch_idx + 1}/{len(self.channels)})   —   {ch}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = C_HEADER_BG

        # ---- Section 2-1 / 2-2 result banner ----
        result_tbl = doc.add_table(rows=3, cols=11)
        result_tbl.style = "Table Grid"

        r0 = result_tbl.rows[0]
        _header_cell(r0.cells[0], "")
        _header_cell(r0.cells[1], "채널명")
        _header_cell(r0.cells[2], "Excitation")
        _header_cell(r0.cells[3], "Gain")
        _header_cell(r0.cells[4], "Offset")
        for j, r in enumerate(rs):
            _header_cell(r0.cells[5 + j], f"{r:.0f}Ω")

        # 2-1 row
        r1 = result_tbl.rows[1]
        _section_cell(r1.cells[0], "2-1\n(100Ω)")
        _para(r1.cells[1], ch, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(r1.cells[2], f"{cal.excitation:.4f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(r1.cells[3], f"{cal.gain:.6f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(r1.cells[4], f"{cal.offset_100:.6f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        for j, r in enumerate(rs):
            dev = cal.dev_final_100.get(r)
            tol = cal.tolerance_max_100.get(r, self.sensor.tolerance_ohm)
            txt = _fmt(dev, 4) if dev is not None else "-"
            c = r1.cells[5 + j]
            _para(c, txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=_pass_color(dev, tol) if dev is not None else None)

        # 2-2 row
        r2 = result_tbl.rows[2]
        _section_cell(r2.cells[0], "2-2\n(Mean)")
        _para(r2.cells[1], ch, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(r2.cells[2], f"{cal.excitation:.4f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(r2.cells[3], f"{cal.gain:.6f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(r2.cells[4], f"{cal.offset_mean:.6f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        for j, r in enumerate(rs):
            dev = cal.dev_final_mean.get(r)
            tol = cal.tolerance_max_mean.get(r, self.sensor.tolerance_ohm)
            txt = _fmt(dev, 4) if dev is not None else "-"
            c = r2.cells[5 + j]
            _para(c, txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=_pass_color(dev, tol) if dev is not None else None)

        for row in result_tbl.rows:
            for cell in row.cells:
                _set_cell_border(cell)

        # ---- Charts (side by side) ----
        doc.add_paragraph()
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        gain_chart   = _make_gain_chart(cal)
        dev_chart    = _make_deviation_chart(cal, rs)

        # We insert two charts into a 1-row 2-col table
        chart_tbl = doc.add_table(rows=1, cols=2)
        chart_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        chart_tbl.rows[0].cells[0].paragraphs[0].add_run().add_picture(gain_chart, width=Cm(8.5))
        chart_tbl.rows[0].cells[1].paragraphs[0].add_run().add_picture(dev_chart,  width=Cm(8.5))

        # ---- Detail statistics table ----
        doc.add_paragraph()
        h2 = doc.add_paragraph("상세 정보 (Detail)")
        h2.runs[0].bold = True
        h2.runs[0].font.size = Pt(10)
        h2.runs[0].font.color.rgb = C_HEADER_BG

        # All detail rows: decimal stats, voltages, resistance before/after gain, offsets
        n_cols = 2 + len(rs)
        dtbl = doc.add_table(rows=0, cols=n_cols)
        dtbl.style = "Table Grid"

        def _add_drow(label_main: str, label_sub: str, values: List, bold=False, bg=None):
            row = dtbl.add_row()
            _para(row.cells[0], label_main, bold=bold, size=8)
            _para(row.cells[1], label_sub, size=8)
            if bg:
                _set_cell_bg(row.cells[0], bg)
                _set_cell_bg(row.cells[1], bg)
            for j, val in enumerate(values):
                txt = _fmt(val, 4) if isinstance(val, float) else str(val) if val is not None else "-"
                _para(row.cells[2 + j], txt, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
                if bg:
                    _set_cell_bg(row.cells[2 + j], bg)

        # Header
        hrow = dtbl.add_row()
        _header_cell(hrow.cells[0], "구분")
        _header_cell(hrow.cells[1], "항목")
        for j, r in enumerate(rs):
            _header_cell(hrow.cells[2 + j], f"{r:.0f}Ω")

        # Decimal
        _add_drow("1) Decimal", "AVG", [int(cal.decimals_avg.get(r, 0)) for r in rs], bg=C_SECTION_BG)
        _add_drow("",           "MIN", [int(cal.decimals_min.get(r, 0)) for r in rs])
        _add_drow("",           "MAX", [int(cal.decimals_max.get(r, 0)) for r in rs])

        # Voltage before gain
        _add_drow("2) 전압(before gain)", "Ref", [cal.voltage_ref.get(r) for r in rs], bg=C_SECTION_BG)
        _add_drow("",                      "AVG", [cal.voltages_avg.get(r) for r in rs])

        # Resistance before gain
        _add_drow("3) 저항(before gain)", "Ref",    [r for r in rs], bg=C_SECTION_BG)
        _add_drow("",                      "AVG",    [cal.r_before_gain.get(r) for r in rs])
        _add_drow("",                      "편차",   [cal.dev_before_gain.get(r) for r in rs])

        # Gain value
        grow = dtbl.add_row()
        _para(grow.cells[0], "4) Gain", bold=True, size=8)
        _para(grow.cells[1], f"{cal.gain:.8f}", size=8)
        _set_cell_bg(grow.cells[0], C_SECTION_BG)
        _set_cell_bg(grow.cells[1], C_SECTION_BG)
        for j in range(len(rs)):
            _set_cell_bg(grow.cells[2 + j], C_SECTION_BG)

        # Resistance after gain
        _add_drow("5) 저항(after gain)",  "AVG",  [cal.r_after_gain_avg.get(r) for r in rs], bg=C_SECTION_BG)
        _add_drow("",                      "MIN",  [cal.r_after_gain_min.get(r) for r in rs])
        _add_drow("",                      "MAX",  [cal.r_after_gain_max.get(r) for r in rs])
        _add_drow("",                      "편차", [cal.dev_after_gain.get(r) for r in rs])

        # Offsets
        orow = dtbl.add_row()
        _para(orow.cells[0], "6) Offset", bold=True, size=8)
        _para(orow.cells[1], f"100Ω: {cal.offset_100:.6f}  /  Mean: {cal.offset_mean:.6f}", size=8)
        _set_cell_bg(orow.cells[0], C_SECTION_BG)
        _set_cell_bg(orow.cells[1], C_SECTION_BG)
        for j in range(len(rs)):
            _set_cell_bg(orow.cells[2 + j], C_SECTION_BG)

        # Method 2-1 final
        _add_drow("7) 2-1 최종(100Ω offset)", "AVG",    [cal.r_final_100_avg.get(r) for r in rs], bg=C_SECTION_BG)
        _add_drow("",                          "편차",   [cal.dev_final_100.get(r) for r in rs])
        _add_drow("",                          "허용(+)", [cal.tolerance_max_100.get(r) for r in rs])
        _add_drow("",                          "허용(-)", [cal.tolerance_min_100.get(r) for r in rs])

        # Method 2-2 final
        _add_drow("8) 2-2 최종(Mean offset)",  "AVG",    [cal.r_final_mean_avg.get(r) for r in rs], bg=C_SECTION_BG)
        _add_drow("",                          "편차",   [cal.dev_final_mean.get(r) for r in rs])
        _add_drow("",                          "허용(+)", [cal.tolerance_max_mean.get(r) for r in rs])
        _add_drow("",                          "허용(-)", [cal.tolerance_min_mean.get(r) for r in rs])

        for row in dtbl.rows:
            for cell in row.cells:
                _set_cell_border(cell)

    # ── Public API ───────────────────────────────────────────────────────────

    def write(self, output_path: str):
        """Generate the full calibration DOCX report and save to output_path."""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = Document()
        self._page_setup(doc)

        # Remove default empty paragraph
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)

        self._add_cover(doc)
        self._add_summary(doc)

        for idx, ch in enumerate(self.channels):
            self._add_channel_page(doc, idx, ch)

        doc.save(output_path)
        return output_path
